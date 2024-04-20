import os
import g4f
import re
import textwrap
import docx

from PyQt6.QtCore import QThread, pyqtSignal
from PyQt6.QtWidgets import (QApplication, QFileDialog, QMainWindow, QPushButton,
                             QTextEdit, QVBoxLayout, QWidget, QProgressBar, QHBoxLayout)

from utils.q_transcriber import QTranscriber
# from downloading_app import DownloadingModelWidget
from utils import functions
import settings

# THEB_AI_API_KEY = os.getenv("THEB_AI_API_KEY")


class TranscribeThread(QThread):
    finished = pyqtSignal()

    def __init__(self, window):
        super().__init__()
        self.window = window

    def run(self):
        task_list = [(Path(self.window.input, fn),
                      Path(self.window.output,
                      Path(fn).stem).with_suffix("." + self.window.args.output_type)
                      ) for fn in os.listdir(self.window.input)
                     ]
        self.window.transcriber.process_task_list_pool(task_list)
        self.finished.emit()


class MainWindow(QMainWindow):
    app_dir = os.path.dirname(os.path.abspath(__file__))

    # def __new__(cls):
    #     dir_path = os.path.join(cls.app_dir, 'models')
    #     if not os.path.exists(dir_path):
    #         os.makedirs(dir_path)
    #     return super().__new__(cls)

    def __init__(self):
        super().__init__()
        self.init_ui()
        self.args, unknown = functions.parse_arguments().parse_known_args()
        self.input = self.get_default_dir("audio")
        self.output = self.get_default_dir("output")
        self.model_dir = settings.model.path_to_model
        self.settings = QSettings('JJoy', 'AppSettings')
        try:
            self.resize(self.settings.value('window size'))
            self.move(self.settings.value('window position'))
            self.input = self.settings.value('input_dir', self.input)
            self.output = self.settings.value('output_dir', self.output)
        except Exception as e:
            print(f"Exception!!! {e}")

    def closeEvent(self, event):
        self.settings.setValue('window size', self.size())
        self.settings.setValue('window position', self.pos())
        self.settings.setValue('input_dir', self.input)
        self.settings.setValue('output_dir',  self.output)
        self.settings.sync()

    def get_default_dir(self, dir_name):
        dir_path = os.path.join(self.app_dir, dir_name)
        if not os.path.exists(dir_path):
            os.makedirs(dir_path)
        return dir_path

    def set_input_dir(self):
        input_dir = QFileDialog.getExistingDirectory(self, 'Выберете папку с аудиофайлами', self.app_dir)
        if input_dir:
            self.input = input_dir
            self.settings.setValue('input_dir', self.input)

    def set_output_dir(self):
        output_dir = QFileDialog.getExistingDirectory(self, 'Выберете папку куда сохранять результат', self.app_dir)
        if output_dir:
            self.output = output_dir
            self.settings.setValue('output_dir', self.output)

    def show_file_dialog(self):
        file_name, _ = QFileDialog.getOpenFileName(self,
                                                   "Выберите файл",
                                                   "",
                                                   "Документы (*.docx *.doc *.txt)",
                                                   f"{self.output}")
        if not file_name:
            return
        doc = docx.Document(file_name)
        text = ''.join(para.text for para in doc.paragraphs)
        self.text_edit.setPlainText(text)

    def save_as_docx(self):
        file_name, _ = QFileDialog.getSaveFileName(self, 'Сохранить в Word', '', 'Word Document (*.docx)')
        if file_name:
            if not file_name.endswith('.docx'):
                file_name = f'{file_name}.docx'
            doc = docx.Document()
            doc.add_paragraph(self.text_edit.toPlainText())
            doc.save(file_name)

    def init_ui(self):
        # Create UI elements
        self.button_transcribe = QPushButton('Транскрибировать аудио', self)
        self.button_transcribe.clicked.connect(self.transcribe_audio)

        self.button_ai = QPushButton('Применить ИИ', self)
        self.button_ai.clicked.connect(self.ai_text_editing)

        self.button_pick_file = QPushButton('Открыть документ', self)
        self.button_pick_file.clicked.connect(self.show_file_dialog)

        self.text_edit = QTextEdit(self)
        self.progress_bar = QProgressBar(self)
        self.progress_bar.setFont(QFont('Arial', 7))

        # Set layout
        layout = QVBoxLayout()
        layout.addWidget(self.button_transcribe)
        self.hbox = QHBoxLayout()
        self.hbox.addWidget(self.button_ai)
        self.hbox.addWidget(self.button_pick_file)
        layout.addLayout(self.hbox)
        layout.addWidget(self.text_edit)
        layout.addWidget(self.progress_bar)

        # Create a central widget and set the layout
        central_widget = QWidget(self)
        central_widget.setLayout(layout)
        self.setCentralWidget(central_widget)

        # Set window title
        icon = QtGui.QIcon()
        icon.addPixmap(QtGui.QPixmap("icons/icon.png"), state=QtGui.QIcon.State.On)
        self.setWindowIcon(icon)
        self.setWindowTitle('JJoy - Audio to Text Converter')
        # self.download_widget = DownloadingModelWidget()
        # self.download_widget.hide()

        # Create menu bar
        menu_bar = self.menuBar()

        # Create 'File' menu
        file_menu = menu_bar.addMenu('Файл')
        settings_menu = menu_bar.addMenu('Настройки')
        info_menu = menu_bar.addMenu('Инфо')

        # Add 'Save' action to the 'File' menu
        save_action = QAction('Сохранить', self)
        save_action.triggered.connect(self.save_as_docx)
        file_menu.addAction(save_action)

        # Add 'Select Model Folder' action to the 'File' menu
        select_model_folder_action = QAction('Выбрать модель', self)
        select_model_folder_action.triggered.connect(self.select_model_folder)
        settings_menu.addAction(select_model_folder_action)

        select_input_folder_action = QAction('Указать папку с аудио', self)
        select_input_folder_action.triggered.connect(self.set_input_dir)
        settings_menu.addAction(select_input_folder_action)

        select_output_folder_action = QAction('Указать папку для сохранения результатов', self)
        select_output_folder_action.triggered.connect(self.set_output_dir)
        settings_menu.addAction(select_output_folder_action)

        about_action = QAction('Требования', self)
        about_action.triggered.connect(self.show_reqs)
        info_menu.addAction(about_action)

        author_action = QAction('Об авторе', self)
        author_action.triggered.connect(self.show_author)
        info_menu.addAction(author_action)


    def show_reqs(self):
        text = f"""Чтобы заработало нужно:<br>
                 1.Скачать <a href={settings.ffmpeg_zip_url_win}>ffmpeg</a> и распаковать в C:\\ffmpeg <br>
                 2. В настройках выбрать папки с аудио и куда сохранять. В папке с аудио должно быть только аудио. 
                 3. Для более точного распознавания нужно скачать <a href={settings.big_model.url}>большую модель</a>
                  и распаковать в директорию с программой. Перейти в настройки и выбрать папку с моделью.<br> 

                  """
        functions.show_interactive_text(self, text)

    def show_author(self):
        text= """
        Разработал Кубышко Павел.<br>
        Использовались модели vosk и g4f.<br>
        Если вам понравилось приложение, отблагодарите автора <br>
         <a href='https://yoomoney.ru/to/410012760894954'>чашечкой кофе</a>.
         """
        functions.show_interactive_text(self, text)

    def select_model_folder(self):
        folder_name = QFileDialog.getExistingDirectory(self, 'Выбрать папку с моделью', '')
        if folder_name:
            self.model_dir = folder_name
            self.save_model_dir(self.model_dir)

    def save_model_dir(self, model_folder):
        # Сохраняем путь в настройках приложения
        settings = QSettings('JJoy', 'AppSettings')
        settings.setValue('model_dir', model_folder)
        self.model_dir = model_folder
        self.create_transcriber()

    def create_transcriber(self):
        self.args.model = self.model_dir
        logging.getLogger().setLevel("INFO")
        try:
            self.transcriber = QTranscriber(self.args)
        except:
            print("нет модели")

    def transcribe_audio(self):
        self.button_transcribe.setText('Идет транскрибация...')
        self.button_transcribe.setEnabled(False)
        self.menuBar().setEnabled(False)
        # QTimer.singleShot(100, self.start_transcribing)  # Add a 100ms delay before starting the transcription
        self.transcribe_thread = TranscribeThread(self)
        self.transcribe_thread.finished.connect(self.transcription_finished)
        self.transcribe_thread.start()

    def transcription_finished(self):
        self.button_transcribe.setText('Транскрибировать аудио')
        self.button_transcribe.setEnabled(True)
        self.menuBar().setEnabled(True)

    def ai_text_editing(self):
        text_to_edit = self.text_edit.toPlainText()
        if not text_to_edit:
            print('Нет текста')
            return
        print(g4f.version)  # check version
        active_providers = [p for p in g4f.Provider.__providers__[13:] if p.working]
        print(active_providers)
        batches = textwrap.wrap(text_to_edit, 1500)
        ai_response = ''
        self.progress_bar.setRange(0, len(batches))

        for i, batch in enumerate(batches):
            for provider in active_providers:
                print(provider)
                try:
                    response = g4f.ChatCompletion.create(
                        model=g4f.models.default,
                        messages=[{"role": "user",
                                   "content": f"Ты получишь транскрибацию аудио файла. Нужно переделать её в удобный для "
                                              f"чтения вид: поставить знаки препинания, исправить грамматику, прямую, "
                                              f"косвенную речь. Старайся как можно меньше изменять смысл текста. "
                                              f"СВОИ КОММЕНТАРИИИ НЕ ПИШИ!!! ОЦЕНОЧНЫХ СУЖДЕНИЙ НЕ ВЫСКАЗЫВАЙ!"
                                              f"Вот текст для исправления: {batch}"}],
                        provider=provider,
                        virtual_display = True,
                        # stream = True
                        # auth=THEB_AI_API_KEY,
                        # timeout=30
                    )  # alternative model setting
                    self.progress_bar.setValue(i + 1)
                except Exception:
                    print('exception')
                    # self.text_edit.setPlainText("Возможно проблемы с интернетом")
                else:
                    if response:
                        ai_response += f'\n\n{response}'
                        break
        ai_response = self.remove_tags(ai_response, ["PHIND_SPAN_BEGIN", "PHIND_SPAN_END"]).strip()
        self.text_edit.setPlainText(ai_response)

    @staticmethod
    def remove_tags(text, tag_names):
        for tag_name in tag_names:
            pattern = re.compile(r'<' + tag_name + r'>.*?</' + tag_name + r'>', re.DOTALL)
            text = re.sub(pattern, '', text)
        return text

    # def check_model_dir(self):
    #     # # Проверяем директорию модели в папке приложения
    #     # app_dir = os.path.dirname(os.path.abspath(__file__))
    #     # model_dir = os.path.join(app_dir, model.dir_name)  # 'vosk-model-ru-0. 42')
    #
    #     if os.path.exists(self.model_dir):
    #         self.save_model_dir(self.model_dir)
    #     else:
    #         # Если нет - предлагаем скачать
    #         resp = QMessageBox.question(
    #             self, 'Model not found',
    #             'Модель для транскрибации не найдена. Скачать?',
    #             QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
    #
    #         if resp == QMessageBox.StandardButton.Yes:
    #             model_zip = model.zip_name
    #
    #             self.download_widget.show()
    #             self.download_widget.download(model)
    #
    #             app_dir = os.path.dirname(os.path.abspath(__file__))
    #             with zipfile.ZipFile(model_zip, 'r') as zip_ref:
    #                 zip_ref.extractall(app_dir)
    #             os.remove(model_zip)
    #             self.save_model_dir(model.dir_name)
    #         else:
    #             # Если отказались - запрашиваем директорию
    #             model_dir = QFileDialog.getExistingDirectory(self, 'Select Model Directory')
    #             print(model_dir)
    #             if model_dir:
    #                 # Сохраняем выбранную директорию
    #                 self.save_model_dir(model_dir)
    #             print(self.model_dir)


if __name__ == "__main__":
    import logging
    import sys
    import os

    from pathlib import Path

    from PyQt6 import QtGui
    from PyQt6.QtCore import QSettings, QRect, QThread, pyqtSignal
    from PyQt6.QtGui import QAction, QFont

    import settings

    # model = settings.model
    app = QApplication(sys.argv)
    screen_width = QApplication.screens()[0].size().width()
    screen_height = QApplication.screens()[0].size().height()
    window = MainWindow()
    win_width = 350
    win_height = 300
    window.setGeometry(QRect((screen_width - win_width) // 2,
                             (screen_height - win_height) // 2,
                             win_width,
                             win_height)
                       )



    # if args.list_models is True:
    #     list_models()

    # if args.list_languages is True:
    #     list_languages()

    # if not args.input:
    #     logging.info("Please specify input file or directory")
    #     args.input = window.set_input_dir()
    #
    # if not Path(args.input).exists():
    #     logging.info("File/folder {args.input} does not exist, " \
    #                  "please specify an existing file/directory")
    #     window.set_output_dir()

    window.show()  # будем хоть какоето окно видеть, пока запускается

    window.create_transcriber()


    sys.exit(app.exec())

#     def main(input, output, log_level='INFO'):
#
#
#         transcriber = Transcriber(args)
#
#         if Path(args.input).is_dir():
#             task_list = [(Path(args.input, fn),
#                           Path(args.output,
#                                Path(fn).stem).with_suffix("." + args.output_type)) for fn in os.listdir(args.input)]
#         elif Path(args.input).is_file():
#             if args.output == "":
#                 task_list = [(Path(args.input), args.output)]
#             else:
#                 task_list = [(Path(args.input), Path(args.output))]
#         else:
#             logging.info("Wrong arguments")
#             sys.exit(1)
#         # print(type(list_models()))
#
#         transcriber.process_task_list_pool(task_list)
#
#
# main("audio", "output")
